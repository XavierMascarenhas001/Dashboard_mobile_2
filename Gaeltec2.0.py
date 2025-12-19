# dashboard_mapped.py
import streamlit as st
import pandas as pd
import plotly.express as px
import re
import geopandas as gpd
import pydeck as pdk
import os
import glob
from PIL import Image
from io import BytesIO
import base64
import plotly.graph_objects as go
import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
from collections import OrderedDict
from difflib import SequenceMatcher
from rapidfuzz import fuzz, process

# --- Page config for wide layout ---
st.set_page_config(
    page_title="Gaeltec Dashboard",
    layout="wide",
    initial_sidebar_state="expanded"
)

def sanitize_sheet_name(name: str) -> str:
    """
    Remove or replace invalid characters for Excel sheet names.
    """
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)
    return name[:31]

def get_scottish_weather(api_key, location="Ayrshire"):
    """
    Get weather data for Scottish locations
    """
    locations = {
        "Ayrshire": {"lat": 55.458, "lon": -4.629},
        "Lanarkshire": {"lat": 55.676, "lon": -3.785},
        "Glasgow": {"lat": 55.864, "lon": -4.252},
        "Edinburgh": {"lat": 55.953, "lon": -3.188}
    }
    
    if location in locations:
        coords = locations[location]
    else:
        coords = locations["Ayrshire"]
    
    base_url = "http://api.openweathermap.org/data/2.5/weather"
    params = {
        'lat': coords["lat"],
        'lon': coords["lon"],
        'appid': api_key,
        'units': 'metric'
    }
    
    try:
        response = requests.get(base_url, params=params, timeout=10)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        st.error(f"Error fetching weather data: {e}")
        return None

@st.cache_data(ttl=1800)
def get_weather_forecast(api_key, location="Ayrshire"):
    """
    Get 5-day forecast for Scottish locations
    """
    locations = {
        "Ayrshire": {"lat": 55.458, "lon": -4.629},
        "Lanarkshire": {"lat": 55.676, "lon": -3.785}
    }
    
    if location in locations:
        coords = locations[location]
    else:
        coords = locations["Ayrshire"]
    
    base_url = "http://api.openweathermap.org/data/2.5/forecast"
    params = {
        'lat': coords["lat"],
        'lon': coords["lon"],
        'appid': api_key,
        'units': 'metric'
    }
    
    try:
        response = requests.get(base_url, params=params, timeout=10)
        response.raise_for_status()
        return response.json()
    except Exception as e:
        st.error(f"Forecast API error: {e}")
        return None

def assess_construction_impact(weather_data):
    """Assess construction impact based on weather conditions"""
    if not weather_data:
        return "No weather data available"
    
    try:
        conditions = {
            'temp': weather_data['main']['temp'],
            'description': weather_data['weather'][0]['description'].lower(),
            'wind_speed': weather_data['wind']['speed'],
            'humidity': weather_data['main']['humidity']
        }
        
        impacts = []
        
        # Temperature impact
        if conditions['temp'] < 0:
            impacts.append("❄️ **High Impact**: Freezing conditions - ground work difficult")
        elif conditions['temp'] < 5:
            impacts.append("🌡️ **Medium Impact**: Cold conditions - slower work pace")
        elif conditions['temp'] > 25:
            impacts.append("🔥 **Medium Impact**: Hot conditions - hydration breaks needed")
        else:
            impacts.append("✅ **Low Impact**: Ideal temperature conditions")
        
        # Precipitation impact
        precip_keywords = ['rain', 'drizzle', 'shower', 'storm', 'thunder']
        if any(keyword in conditions['description'] for keyword in precip_keywords):
            impacts.append("🌧️ **High Impact**: Wet conditions - outdoor work limited")
        
        # Wind impact
        if conditions['wind_speed'] > 10:
            impacts.append("💨 **High Impact**: High winds - unsafe for work at height")
        elif conditions['wind_speed'] > 6:
            impacts.append("💨 **Medium Impact**: Windy conditions - use caution")
        
        # Humidity impact
        if conditions['humidity'] > 80:
            impacts.append("💧 **Low Impact**: High humidity - equipment may be affected")
        
        return "\n\n".join(impacts) if impacts else "✅ **Low Impact**: Favorable conditions for construction"
    
    except Exception as e:
        return f"Could not assess impact: {str(e)}"

def poles_to_word(df: pd.DataFrame) -> BytesIO:
    doc = Document()

    df = df.copy()
    df = df.replace(to_replace=["nan", "NaN", "None", None], value="")

    grouped = df.groupby('pole', sort=False)

    for pole, group in grouped:
        pole_str = str(pole).strip()
        if not pole_str:
            continue

        unique_texts = OrderedDict()

        for _, row in group.iterrows():
            parts = []
            wi = str(row.get('Work instructions', '')).strip()
            comment = str(row.get('comment', '')).strip()

            if wi:
                parts.append(wi)
            if comment:
                parts.append(f"({comment})")
            
            if parts:
                text = " ".join(parts)
                normalized = text.lower().strip()
                unique_texts[normalized] = text

        if not unique_texts:
            continue

        p = doc.add_paragraph(style='List Bullet')
        run_number = p.add_run(f"{pole_str} – ")
        run_number.bold = True
        run_number.font.name = 'Times New Roman'
        run_number.font.size = Pt(12)

        texts = list(unique_texts.values())
        for i, text in enumerate(texts):
            run_item = p.add_run(text)
            run_item.bold = True
            run_item.font.name = 'Times New Roman'
            run_item.font.size = Pt(12)

            if "Erect Pole" in text:
                run_item.font.highlight_color = WD_COLOR_INDEX.RED

            if i < len(texts) - 1:
                p.add_run(" ; ")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def normalize_columns(df):
    """Normalize column names to lowercase with underscores"""
    df = df.copy()
    df.columns = (
        df.columns
        .astype(str)
        .str.strip()
        .str.lower()
        .str.replace(" ", "_")
        .str.replace(r"[^\w]", "_", regex=True)
    )
    return df

@st.cache_data
def load_data_files():
    """Load all required data files"""
    files_to_load = {
        "aggregated": "CF_aggregated.parquet",
        "resume": "CF_resume.parquet",
        "miscellaneous": "miscelaneous.parquet",
        "pid": "Resume_PID.parquet"
    }
    
    data = {}
    
    for name, filepath in files_to_load.items():
        try:
            if os.path.exists(filepath):
                df = pd.read_parquet(filepath)
                df = normalize_columns(df)
                data[name] = df
                st.sidebar.success(f"✓ {name.replace('_', ' ').title()} loaded")
            else:
                st.sidebar.warning(f"⚠ {filepath} not found")
                data[name] = None
        except Exception as e:
            st.sidebar.error(f"✗ Failed to load {filepath}: {e}")
            data[name] = None
    
    return data

    
# --- MAPPINGS ---

# --- Project Manager Mapping ---
project_mapping = {
    "Jonathon Mcclung": ["Ayrshire", "PCB"],
    "Gary MacDonald": ["Ayrshire", "LV"],
    "Jim Gaffney": ["Lanark", "PCB"],
    "Calum Thomson": ["Ayrshire", "Connections"],
    "Calum Thomsom": ["Ayrshire", "Connections"],
    "Calum Thompson": ["Ayrshire", "Connections"],
    "Andrew Galt": ["Ayrshire", "-"],
    "Henry Gordon": ["Ayrshire", "-"],
    "Jonathan Douglas": ["Ayrshire", "11 kV"],
    "Jonathon Douglas": ["Ayrshire", "11 kV"],
    "Matt": ["Lanark", ""],
    "Lee Fraser": ["Ayrshire", "Connections"],
    "Lee Frazer": ["Ayrshire", "Connections"],
    "Mark": ["Lanark", "Connections"],
    "Mark Nicholls": ["Ayrshire", "Connections"],
    "Cameron Fleming": ["Lanark", "Connections"],
    "Ronnie Goodwin": ["Lanark", "Connections"],
    "Ian Young": ["Ayrshire", "Connections"],
    "Matthew Watson": ["Lanark", "Connections"],
    "Aileen Brese": ["Ayrshire", "Connections"],
    "Mark McGoldrick": ["Lanark", "Connections"]
}

# --- Region Mapping ---
mapping_region = {
    "Newmilns": ["Irvine Valley"],
    "New Cumnock": ["New Cumnock"],
    "Kilwinning": ["Kilwinning"],
    "Stewarton": ["Irvine Valley"],
    "Kilbirnie": ["Kilbirnie and Beith"],
    "Coylton": ["Ayr East"],
    "Irvine": ["Irvine Valley", "Irvine East", "Irvine West"],
    "TROON": ["Troon"],
    "Ayr": ["Ayr East", "Ayr North", "Ayr West"],
    "Maybole": ["Maybole, North Carrick and Coylton"],
    "Clerkland": ["Irvine Valley"],
    "Glengarnock": ["Kilbirnie and Beith"],
    "Ayrshire": ["North Coast and Cumbraes","Prestwick", "Saltcoats and Stevenston", "Troon", "Ayr East", "Ayr North",
                 "Ayr West","Annick","Ardrossan and Arran","Dalry and West Kilbride","Girvan and South Carrick","Irvine East",
                 "Irvine Valley","Irvine West","Kilbirnie and Beith","Kilmarnock East and Hurlford","Kilmarnock North",
                 "Kilmarnock South","Kilmarnock West and Crosshouse","Kilwinning","Kyle","Maybole, North Carrick and Coylton",
                 "Ayr, Carrick and Cumnock","East_Ayrshire","North_Ayrshre","South_Ayrshre","Doon Valley"],
    "Lanark": ["Abronhill, Kildrum and the Village","Airdrie Central","Airdrie North","Airdrie South","Avondale and Stonehouse",
               "Ballochmyle","Bellshill","Blantyre","Bothwell and Uddingston","Cambuslang East","Cambuslang West",
               "Clydesdale East","Clydesdale North","Clydesdale South","Clydesdale West","Coatbridge North and Glenboig",
               "Coatbridge South","Coatbridge West","Cumbernauld North","Cumbernauld South",
               "East Kilbride Central North","East Kilbride Central South","East Kilbride East","East Kilbride South",
               "East Kilbride West","Fortissat","Hamilton North and East","Hamilton South","Hamilton West and Earnock",
               "Mossend and Holytown","Motherwell North","Motherwell South East and Ravenscraig","Motherwell West",
               "Rutherglen Central and North","Rutherglen South","Strathkelvin","Thorniewood","Wishaw","Larkhall",
               "Airdrie and Shotts","Cumbernauld, Kilsyth and Kirkintilloch East","East Kilbride, Strathaven and Lesmahagow",
               "Lanark and Hamilton East","Motherwell and Wishaw","North_Lanarkshire","South_Lanarkshire"]
}

# --- File Project Mapping ---
file_project_mapping = {
    "pcb 2022": ["Ayrshire", "PCB"],
    "33kv refurb": ["Ayrshire", "33kv Refurb"],
    "connections": ["Ayrshire", "Connections"],
    "storms": ["Ayrshire", "Storms"],
    "11kv refurb": ["Ayrshire", "11kv Refurb"],
    "aurs road": ["Ayrshire", "Aurs Road"],
    "spen labour": ["Ayrshire", "SPEN Labour"],
    "lvhi5": ["Ayrshire", "LV"],
    "pcb": ["Ayrshire", "PCB"],
    "lanark": ["Lanark", ""],
    "11kv refur": ["Lanark", "11kv Refurb"],
    "lv & esqcr": ["Lanark", "LV"],
    "11kv rebuilt": ["Lanark", "11kV Rebuilt"],
    "33kv rebuilt": ["Lanark", "33kV Rebuilt"]
}

# --- Pole Mappings (dictionary style, includes new additions) ---
pole_keys = {
    "9x220 BIOCIDE LV POLE": "9m B",
    "9x275 BIOCIDE LV POLE": "9s B",
    "9x220 CREOSOTE LV POLE": "9m",
    "9x275 CREOSOTE LV POLE": "9s",
    "9x220 HV SINGLE POLE": "9m",
    "9x275 HV SINGLE POLE": "9s",
    "9x295 HV SINGLE POLE": "9es",
    "9x315 HV SINGLE POLE": "9esp",
    "10x230 BIOCIDE LV POLE": "10m B",
    "10x230 HV SINGLE POLE": "10m",
    "10x285 BIOCIDE LV POLE": "10s B",
    "10x285 H POLE HV Creosote": "10s",
    "10x285 HV SINGLE POLE": "10s",
    "10x305 HV SINGLE POLE": "10es",
    "11x295 HV SINGLE POLE": "11s",
    "11x295 H POLE HV Creosote": "11s",
    "11x295 BIOCIDE LV POLE": "11sB",
    "12x250 BIOCIDE LV POLE": "12m B",
    "12x305 BIOCIDE LV POLE": "12s B",
    "12x250 CREOSOTE LV POLE": "12m",
    "12x305 CREOSOTE LV POLE": "12s",
    "12x305 H POLE HV Creosote":"12s",
    "12x250 HV SINGLE POLE": "12m",
    "12x305 HV SINGLE POLE": "12s",
    "12x325 HV SINGLE POLE": "12es",
    "12x345 HV SINGLE POLE": "12esp",
    "13x260 BIOCIDE LV POLE": "13m B",
    "13x320 BIOCIDE LV POLE": "13s B",
    "13x260 CREOSOTE LV POLE": "13m",
    "13x320 CREOSOTE LV POLE": "13s",
    "13x260 HV SINGLE POLE": "13m",
    "13x320 HV SINGLE POLE": "13s",
    "13x340 HV SINGLE POLE": "13es",
    "13x365 HV SINGLE POLE": "13esp",
    "14x275 BIOCIDE LV POLE": "14m B",
    "14x335 BIOCIDE LV POLE": "14s B",
    "14x275 CREOSOTE LV POLE": "14m",
    "14x335 CREOSOTE LV POLE": "14s",
    "14x275 HV SINGLE POLE": "14m",
    "14x335 HV SINGLE POLE": "14s",
    "14x355 HV SINGLE POLE": "14es",
    "14x375 HV SINGLE POLE": "14esp",
    "16x305 BIOCIDE LV POLE": "16m B",
    "16x365 BIOCIDE LV POLE": "16s B",
    "16x305 CREOSOTE LV POLE": "16m",
    "16x365 CREOSOTE LV POLE": "16s",
    "16x305 HV SINGLE POLE": "16m",
    "16x365 HV SINGLE POLE": "16s",
    "16x385 HV SINGLE POLE": "16es",
    "16x405 HV SINGLE POLE": "16esp",
}

# --- Equipment / Conductor Mappings ---
equipment_keys = {
    "Hazel - 50mm² AAAC bare (1000m drums)": "Hazel 50mm²",
    "Oak - 100mm² AAAC bare (1000m drums)": "Oak 100mm²",
    "Ash - 150mm² AAAC bare (1000m drums)": "Ash 150mm²",
    "Poplar - 200mm² AAAC bare (1000m drums)": "Poplar 200mm²",
    "Upas - 300mm² AAAC bare (1000m drums)": "Upas 300mm²",
    "Poplar OPPC - 200mm² AAAC equivalent bare": "Poplar OPPC 200mm²",
    "Upas OPPC - 300mm² AAAC equivalent bare": "Upas OPPC 300mm²",
    # ACSR
    "Gopher - 25mm² ACSR bare (1000m drums)": "Gopher 25mm²",
    "Caton - 25mm² Compacted ACSR bare (1000m drums)": "Caton 25mm²",
    "Rabbit - 50mm² ACSR bare (1000m drums)": "Rabbit 50mm²",
    "Wolf - 150mm² ACSR bare (1000m drums)": "Wolf 150mm²",
    "Horse - 70mm² ACSR bare": "Horse 70mm²",
    "Dog - 100mm² ACSR bare (1000m drums)": "Dog 100mm²",
    "Dingo - 150mm² ACSR bare (1000m drums)": "Dingo 150mm²",
    # Copper
    "Hard Drawn Copper 16mm² ( 3/2.65mm ) (500m drums)": "Copper 16mm²",
    "Hard Drawn Copper 32mm² ( 3/3.75mm ) (1000m drums)": "Copper 32mm²",
    "Hard Drawn Copper 70mm² (500m drums)": "Copper 70mm²",
    "Hard Drawn Copper 100mm² (500m drums)": "Copper 100mm²",
    # PVC covered
    "35mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 35mm² GY PVC",
    "70mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 70mm² GY PVC",
    "35mm² Copper (Blue PVC covered) (50m drums)": "Copper 35mm² Blue PVC",
    "70mm² Copper (Blue PVC covered) (50m drums)": "Copper 70mm² Blue PVC",
    # Double insulated
    "35mm² Double Insulated (Brown) (50m drums)": "Double Insulated 35mm² Brown",
    "35mm² Double Insulated (Blue) (50m drums)": "Double Insulated 35mm² Blue",
    "70mm² Double Insulated (Brown) (50m drums)": "Double Insulated 70mm² Brown",
    "70mm² Double Insulated (Blue) (50m drums)": "Double Insulated 70mm² Blue",
    "120mm² Double Insulated (Brown) (50m drums)": "Double Insulated 120mm² Brown",
    "120mm² Double Insulated (Blue) (50m drums)": "Double Insulated 120mm² Blue",
    # LV cables
    "LV Cable 1ph 4mm Concentric (250m drums)": "LV 1ph 4mm Concentric",
    "LV Cable 1ph 25mm CNE (250m drums)": "LV 1ph 25mm CNE",
    "LV Cable 1ph 25mm SNE (100m drums)": "LV 1ph 25mm SNE",
    "LV Cable 1ph 35mm CNE (250m drums)": "LV 1ph 35mm CNE",
    "LV Cable 1ph 35mm SNE (100m drums)": "LV 1ph 35mm SNE",
    "LV Cable 3ph 35mm Cu Split Con (250m drums)": "LV 3ph 35mm Cu Split Con",
    "LV Cable 3ph 35mm SNE (250m drums)": "LV 3ph 35mm SNE",
    "LV Cable 3ph 35mm CNE (250m drums)": "LV 3ph 35mm CNE",
    "LV Cable 3ph 35mm CNE Al (LSOH) (250m drums)": "LV 3ph 35mm CNE Al LSOH",
    "LV Cable 3c 95mm W/F (250m drums)": "LV 3c 95mm W/F",
    "LV Cable 3c 185mm W/F (250m drums)": "LV 3c 185mm W/F",
    "LV Cable 3c 300mm W/F (250m drums)": "LV 3c 300mm W/F",
    "LV Cable 4c 95mm W/F (250m drums)": "LV 4c 95mm W/F",
    "LV Cable 4c 185mm W/F (250m drums)": "LV 4c 185mm W/F",
    "LV Cable 4c 240mm W/F (250m drums)": "LV 4c 240mm W/F",
    "LV Marker Tape (365m roll)": "LV Marker Tape",
    # 11kV
    "11kv Cable 95mm 3c Poly (250m drums)": "11kV 3c 95mm Poly",
    "11kv Cable 185mm 3c Poly (250m drums)": "11kV 3c 185mm Poly",
    "11kv Cable 300mm 3c Poly (250m drums)": "11kV 3c 300mm Poly",
    "11kv Cable 95mm 1c Poly (250m drums)": "11kV 1c 95mm Poly",
    "11kv Cable 185mm 1c Poly (250m drums)": "11kV 1c 185mm Poly",
    "11kv Cable 300mm 1c Poly (250m drums)": "11kV 1c 300mm Poly",
    "11kV Marker Tape (40m roll)": "11kV Marker Tape"
}

# --- Transformer Mappings ---
transformer_keys = {
    "Transformer 1ph 50kVA": "TX 1ph (50kVA)",
    "Transformer 3ph 50kVA": "TX 3ph (50kVA)",
    "Transformer 1ph 100kVA": "TX 1ph (100kVA)",
    "Transformer 1ph 25kVA": "TX 1ph (25kVA)",
    "Transformer 3ph 200kVA": "TX 3ph (200kVA)",
    "Transformer 3ph 100kVA": "TX 3ph (100kVA)"
}

# --- Equipment / Conductor Mappings ---
conductor_keys = {
    "Hazel - 50mm² AAAC bare (1000m drums)": "Hazel 50mm²",
    "Oak - 100mm² AAAC bare (1000m drums)": "Oak 100mm²",
    "Ash - 150mm² AAAC bare (1000m drums)": "Ash 150mm²",
    "Poplar - 200mm² AAAC bare (1000m drums)": "Poplar 200mm²",
    "Upas - 300mm² AAAC bare (1000m drums)": "Upas 300mm²",
    "Poplar OPPC - 200mm² AAAC equivalent bare": "Poplar OPPC 200mm²",
    "Upas OPPC - 300mm² AAAC equivalent bare": "Upas OPPC 300mm²",
    # ACSR
    "Gopher - 25mm² ACSR bare (1000m drums)": "Gopher 25mm²",
    "Caton - 25mm² Compacted ACSR bare (1000m drums)": "Caton 25mm²",
    "Rabbit - 50mm² ACSR bare (1000m drums)": "Rabbit 50mm²",
    "Wolf - 150mm² ACSR bare (1000m drums)": "Wolf 150mm²",
    "Horse - 70mm² ACSR bare": "Horse 70mm²",
    "Dog - 100mm² ACSR bare (1000m drums)": "Dog 100mm²",
    "Dingo - 150mm² ACSR bare (1000m drums)": "Dingo 150mm²",
    # Copper
    "Hard Drawn Copper 16mm² ( 3/2.65mm ) (500m drums)": "Copper 16mm²",
    "Hard Drawn Copper 32mm² ( 3/3.75mm ) (1000m drums)": "Copper 32mm²",
    "Hard Drawn Copper 70mm² (500m drums)": "Copper 70mm²",
    "Hard Drawn Copper 100mm² (500m drums)": "Copper 100mm²",
    # PVC covered
    "35mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 35mm² GY PVC",
    "70mm² Copper (Green / Yellow PVC covered) (50m drums)": "Copper 70mm² GY PVC",
    "35mm² Copper (Blue PVC covered) (50m drums)": "Copper 35mm² Blue PVC",
    "70mm² Copper (Blue PVC covered) (50m drums)": "Copper 70mm² Blue PVC",
    # Double insulated
    "35mm² Double Insulated (Brown) (50m drums)": "Double Insulated 35mm² Brown",
    "35mm² Double Insulated (Blue) (50m drums)": "Double Insulated 35mm² Blue",
    "70mm² Double Insulated (Brown) (50m drums)": "Double Insulated 70mm² Brown",
    "70mm² Double Insulated (Blue) (50m drums)": "Double Insulated 70mm² Blue",
    "120mm² Double Insulated (Brown) (50m drums)": "Double Insulated 120mm² Brown",
    "120mm² Double Insulated (Blue) (50m drums)": "Double Insulated 120mm² Blue"
}

    # LV cables per meter
conductor_2_keys = {
    "LV Cable 1ph 4mm Concentric (250m drums)": "LV 1ph 4mm Concentric",
    "LV Cable 1ph 25mm CNE (250m drums)": "LV 1ph 25mm CNE",
    "LV Cable 1ph 25mm SNE (100m drums)": "LV 1ph 25mm SNE",
    "LV Cable 1ph 35mm CNE (250m drums)": "LV 1ph 35mm CNE",
    "LV Cable 1ph 35mm SNE (100m drums)": "LV 1ph 35mm SNE",
    "LV Cable 3ph 35mm Cu Split Con (250m drums)": "LV 3ph 35mm Cu Split Con",
    "LV Cable 3ph 35mm SNE (250m drums)": "LV 3ph 35mm SNE",
    "LV Cable 3ph 35mm CNE (250m drums)": "LV 3ph 35mm CNE",
    "LV Cable 3ph 35mm CNE Al (LSOH) (250m drums)": "LV 3ph 35mm CNE Al LSOH",
    "LV Cable 3c 95mm W/F (250m drums)": "LV 3c 95mm W/F",
    "LV Cable 3c 185mm W/F (250m drums)": "LV 3c 185mm W/F",
    "LV Cable 3c 300mm W/F (250m drums)": "LV 3c 300mm W/F",
    "LV Cable 4c 95mm W/F (250m drums)": "LV 4c 95mm W/F",
    "LV Cable 4c 185mm W/F (250m drums)": "LV 4c 185mm W/F",
    "LV Cable 4c 240mm W/F (250m drums)": "LV 4c 240mm W/F",
    "LV Marker Tape (365m roll)": "LV Marker Tape",
    # 11kV
    "11kv Cable 95mm 3c Poly (250m drums)": "11kV 3c 95mm Poly",
    "11kv Cable 185mm 3c Poly (250m drums)": "11kV 3c 185mm Poly",
    "11kv Cable 300mm 3c Poly (250m drums)": "11kV 3c 300mm Poly",
    "11kv Cable 95mm 1c Poly (250m drums)": "11kV 1c 95mm Poly",
    "11kv Cable 185mm 1c Poly (250m drums)": "11kV 1c 185mm Poly",
    "11kv Cable 300mm 1c Poly (250m drums)": "11kV 1c 300mm Poly",
    "11kV Marker Tape (40m roll)": "11kV Marker Tape"
}


equipment_keys = {
    "Noja": "Noja",
    "11kV PMSW (Soule)": "11kV PMSW (Soule)",
    "11kv ABSW Hookstick Standard": "11kv ABSW Hookstick Standard",
    "11kv ABSW Hookstick Spring loaded mech": "11kv ABSW Hookstick Spring loaded mech",
    "33kv ABSW Hookstick Dependant": "33kv ABSW Hookstick Dependant",
    "100A LV Fuse JPU 82.5mm": "100A LV Fuse JPU 82.5mm",
    "160A LV Fuse JPU 82.5mm": "160A LV Fuse JPU 82.5mm",
    "200A LV Fuse JPU 82.5mm": "200A LV Fuse JPU 82.5mm",
    "315A LV Fuse JPU 82.5mm": "315A LV Fuse JPU 82.5mm",
    "400A LV Fuse JPU 82.5mm": "400A LV Fuse JPU 82.5mm",
    "200A LV Fuse JSU 92mm": "200A LV Fuse JSU 92mm",
    "315A LV Fuse JSU 92mm": "315A LV Fuse JSU 92mm",
    "400A LV Fuse JSU 92mm": "400A LV Fuse JSU 92mm",
    "100A LV Fuse - Porcelain screw-in": "100A LV Fuse - Porcelain screw-in",
    "160A LV Fuse - Porcelain screw-in": "160A LV Fuse - Porcelain screw-in",
    "200A LV Fuse - Porcelain screw-in": "200A LV Fuse - Porcelain screw-in",
    "Single Phase cut out kit 100A Henley Series 7": "Single Phase cut out kit 100A Henley Series 7",
    "Single Phase SNE Sealing Chamber": "Single Phase SNE Sealing Chamber",
    "Three Phase cut out kit 100A Henley Series 7": "Three Phase cut out kit 100A Henley Series 7",
    "Three Phase 200A Cut out": "Three Phase 200A Cut out",
    "Earth Connector Block 100A 5 Way": "Earth Connector Block 100A 5 Way",
    "Cut out Fuse (MF) 60A": "Cut out Fuse (MF) 60A",
    "Cut out Fuse (MF) 80A": "Cut out Fuse (MF) 80A",
    "Cut out Fuse (MF) 100A": "Cut out Fuse (MF) 100A",
    "Temporary Meter seal white plastic (100)": "Temporary Meter seal white plastic (100)",
    "Meter seals for use with sealing pliers (100)": "Meter seals for use with sealing pliers (100)",
    "Meter sealing wire 200mm long (each)": "Meter sealing wire 200mm long (each)",
    "ABC 1PH & 3PH TERM BOX": "ABC 1PH & 3PH TERM BOX",
    "SINGLE PHASE FUSED ABC BOX": "SINGLE PHASE FUSED ABC BOX",
    "1PH & 3PH FUSED SERV WALL BOX": "1PH & 3PH FUSED SERV WALL BOX",
    "25mm Galvanised Conduit": "25mm Galvanised Conduit",
    "25mm Galvanised Conduit saddles": "25mm Galvanised Conduit saddles",
    "Street Lighting Cut out CNE": "Street Lighting Cut out CNE",
    "Street Lighting Cut out SNE": "Street Lighting Cut out SNE",
    "11KV FUSE UNIT - C-TYPE": "11KV FUSE UNIT - C-TYPE",
    "11KV SOLID LINK - C-TYPE": "11KV SOLID LINK - C-TYPE",
    "11KV OHL ASL C-TYPE RESET 20A 2 SHOT": "11KV OHL ASL C-TYPE RESET 20A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 25A 2 SHOT": "11KV OHL ASL C-TYPE RESET 25A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 1 SHOT": "11KV OHL ASL C-TYPE RESET 40A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 40A 2 SHOT": "11KV OHL ASL C-TYPE RESET 40A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 1 SHOT": "11KV OHL ASL C-TYPE RESET 63A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 2 SHOT": "11KV OHL ASL C-TYPE RESET 63A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 63A 3 SHOT": "11KV OHL ASL C-TYPE RESET 63A 3 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 1 SHOT": "11KV OHL ASL C-TYPE RESET 100A 1 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 2 SHOT": "11KV OHL ASL C-TYPE RESET 100A 2 SHOT",
    "11KV OHL ASL C-TYPE RESET 100A 3 SHOT": "11KV OHL ASL C-TYPE RESET 100A 3 SHOT",
    "11KV FUSE CARRIER - C-TYPE": "11KV FUSE CARRIER - C-TYPE",
    "11KV OHL FUSE ELEMENT C-TYPE 15A": "11KV OHL FUSE ELEMENT C-TYPE 15A",
    "11KV OHL FUSE ELEMENT C-TYPE 25A": "11KV OHL FUSE ELEMENT C-TYPE 25A",
    "11KV OHL FUSE ELEMENT C-TYPE 30A": "11KV OHL FUSE ELEMENT C-TYPE 30A",
    "11KV OHL FUSE ELEMENT C-TYPE 40A": "11KV OHL FUSE ELEMENT C-TYPE 40A",
    "11KV OHL FUSE ELEMENT C-TYPE 50A": "11KV OHL FUSE ELEMENT C-TYPE 50A",
    "11KV OHL ASL - CHEMICAL ACTUATOR": "11KV OHL ASL - CHEMICAL ACTUATOR",
    "11KV OHL ASL DJP-TYPE 20A 2 SHOT": "11KV OHL ASL DJP-TYPE 20A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 1 SHOT": "11KV OHL ASL DJP-TYPE 25A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 25A 2 SHOT": "11KV OHL ASL DJP-TYPE 25A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 1 SHOT": "11KV OHL ASL DJP-TYPE 40A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 40A 2 SHOT": "11KV OHL ASL DJP-TYPE 40A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 1 SHOT": "11KV OHL ASL DJP-TYPE 63A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 2 SHOT": "11KV OHL ASL DJP-TYPE 63A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 63A 3 SHOT": "11KV OHL ASL DJP-TYPE 63A 3 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 1 SHOT": "11KV OHL ASL DJP-TYPE 100A 1 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 2 SHOT": "11KV OHL ASL DJP-TYPE 100A 2 SHOT",
    "11KV OHL ASL DJP-TYPE 100A 3 SHOT": "11KV OHL ASL DJP-TYPE 100A 3 SHOT",
    "11KV OHL FUSE ELEMENT DJP-TYPE 15A": "11KV OHL FUSE ELEMENT DJP-TYPE 15A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 25A": "11KV OHL FUSE ELEMENT DJP-TYPE 25A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 30A": "11KV OHL FUSE ELEMENT DJP-TYPE 30A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 40A": "11KV OHL FUSE ELEMENT DJP-TYPE 40A",
    "11KV OHL FUSE ELEMENT DJP-TYPE 50A": "11KV OHL FUSE ELEMENT DJP-TYPE 50A",
    "0.5 kVa Tx for Noja": "0.5 kVa Tx for Noja",
    "Military Cable for Noja": "Military Cable for Noja",
    "Antenna for Soule or Noja": "Antenna for Soule or Noja",
    "Bracket for antenna": "Bracket for antenna",
    "Coax cable (5m)": "Coax cable (5m)",
    "Antenna for Soule or Noja": "Antenna for Soule or Noja",
    "Bracket for antenna": "Bracket for antenna",
    "Coax cable (5m)": "Coax cable (5m)",
}

insulator_keys = {
    "11kV Pin Insulator; Polymeric": "11kV Pin Insulator; Polymeric",
    "11kV Pin Insulator; Polymeric; High Creepage": "11kV Pin Insulator; Polymeric; High Creepage",
    "33kV Pin Insulator; Porcelain": "33kV Pin Insulator; Porcelain",
    "33kV Post Insulator; Polymeric; Clamp Top Plate": "33kV Post Insulator; Polymeric; Clamp Top Plate",
    "36kV Composite Post Groove Top": "36kV Composite Post Groove Top",
    "11kV Tension Insulator; Polymeric (70kN)": "11kV Tension Insulator; Polymeric (70kN)",
    "33kV Tension Insulator; Polymeric (70kN)": "33kV Tension Insulator; Polymeric (70kN)",
    "36kV Composite Tension Ball/Socket Fitting (125 kN)": "36kV Composite Tension Ball/Socket Fitting (125 kN)",
    "LV / 11kV Stay Insulator": "LV / 11kV Stay Insulator",
    "33kV Stay Insulator": "33kV Stay Insulator",
    "LV Insulator Bobbin Type": "LV Insulator Bobbin Type",
    "LV Insulator Coachscrew Type": "LV Insulator Coachscrew Type"
}


lv_joint_kit_keys = {
    "LVKIT/001": "LVKIT/001 Straight Jt Kit 35mm 1ph CNE/SNE Plastic",
    "LVKIT/002": "LVKIT/002 Straight Jt Kit 35mm 1ph CNE/SNE Pilc",
    "LVKIT/003": "LVKIT/003 Straight Jt Kit 35mm 3ph CNE/SNE Plastic",
    "LVKIT/004": "LVKIT/004 Staight Jt 3ph 35mm XLPE to 4-35 PILC",
    "LVKIT/005": "LVKIT/005 LV Service Cable Stop End",
    "LVKIT/006": "LVKIT/006 LV Service off a service 4-35mm 1/3 phase CNE/SNE",
    "LVKIT/007": "LVKIT/007 LV Service off a service 4-35mm PILC 1ph CNE/SNE",
    "LVKIT/008": "LVKIT/008 Service Pole Term to OHL 1PH CNE",
    "LVKIT/009": "LVKIT/009 Service Pole Term to OHL 1PH SNE",
    "LVKIT/010": "LVKIT/010 Service Pole Term to OHL 3PH 35mm",
    "LVKIT/011": "LVKIT/011 Service Pole Term to Fuses 1PH CNE",
    "LVKIT/012": "LVKIT/012 Service Pole Term to Fuses 1PH SNE",
    "LVKIT/013": "LVKIT/013 Service Pole Term to Fuses 3PH 35mm",
    "LVKIT/014": "LVKIT/014 Service Breech Joint 70-185mm 3c W/F - CNE/SNE",
    "LVKIT/015": "LVKIT/015 Service Breech Joint 240-300mm 3c W/F - CNE/SNE",
    "LVKIT/016": "LVKIT/016 Service Breech Joint 50-95mm PILC - CNE/SNE",
    "LVKIT/017": "LVKIT/017 Service Breech Joint 95-185mm PILC - CNE/SNE",
    "LVKIT/018": "LVKIT/018 Service Breech Joint 185-300mm PILC - CNE/SNE",
    "LVKIT/019": "LVKIT/019 Straight Joint up to 95mm 3c W/F / PILC",
    "LVKIT/020": "LVKIT/020 Straight Joint 185mm 3c W/F / PILC / CONSAC",
    "LVKIT/021": "LVKIT/021 Straight Joint 300mm 3c W/F / PILC / CONSAC",
    "LVKIT/022": "LVKIT/022 Mains Breech Joint 70-95mm 3c W/F",
    "LVKIT/023": "LVKIT/023 Mains Breech Joint 185mm 3c W/F",
    "LVKIT/024": "LVKIT/024 Mains Breech Joint 240/300mm 3c W/F",
    "LVKIT/025": "LVKIT/025 Mains Breech Joint 70-95mm W/F / 50-95mm PILC",
    "LVKIT/026": "LVKIT/026 Mains Breech Joint 185mm W/F / 95-185mm PILC",
    "LVKIT/027": "LVKIT/027 Mains Breech Joint 240/300mm W/F / 185-300mm PILC",
    "LVKIT/028": "LVKIT/028 Loop / V Joint 50-95mm W/F / PILC",
    "LVKIT/029": "LVKIT/029 Loop / V Joint >95-300mm W/F / PILC",
    "LVKIT/030": "LVKIT/030 Y / 3 Loose end Joint 50-185mm W/F / PILC / Districable",
    "LVKIT/031": "LVKIT/031 Y / 3 Loose end Joint 185-300mm W/F / PILC / Districable",
    "LVKIT/032": "LVKIT/032 Stop End 70-95mm W/F / CONSAC",
    "LVKIT/033": "LVKIT/033 Stop End 185-300mm W/F / CONSAC",
    "LVKIT/034": "LVKIT/034 Stop End 50-95mm PILC",
    "LVKIT/035": "LVKIT/035 Stop End 95-300mm PILC",
    "LVKIT/037": "LVKIT/037 Pole Term to OHL 70-95mm W/F",
    "LVKIT/038": "LVKIT/038 Pole Term to OHL 185mm W/F",
    "LVKIT/039": "LVKIT/039 Pole Term to Fuses 70-95mm W/F",
    "LVKIT/040": "LVKIT/040 Pole Term to Fuses 185mm W/F"
}


lv_joint_module_keys = {
    "LVMOD/001": "LVMOD/001 Armour bond module for PILC Service cable Stop Ends",
    "LVMOD/002": "LVMOD/002 Branch connector module for service cables",
    "LVMOD/003": "LVMOD/003 Phase connector remake module for service cables",
    "LVMOD/004": "LVMOD/004 XL Brass tunnel connector module for old PILC concentric cables",
    "LVMOD/005": "LVMOD/005 Insulated insulating piercing mains/service branch connector module (up to 185mm2)",
    "LVMOD/006": "LVMOD/006 Insulated insulating piercing mains/service branch connector module (240-300mm2)",
    "LVMOD/007": "LVMOD/007 Brass neutral earth connector module",
    "LVMOD/008": "LVMOD/008 CONSAC Brass neutral earth connector module",
    "LVMOD/009": "LVMOD/009 95mm2 straight type channel connector module",
    "LVMOD/011": "LVMOD/011 185mm2 straight type channel connector module",
    "LVMOD/013": "LVMOD/013 300mm2 straight type channel connector module",
    "LVMOD/015": "LVMOD/015 95mm2 branch type channel connector module",
    "LVMOD/017": "LVMOD/017 185mm2 branch type channel connector module",
    "LVMOD/018": "LVMOD/018 185mm2 branch type channel connector c/w brass non-shear bolts module",
    "LVMOD/019": "LVMOD/019 300mm2 branch type channel connector module",
    "LVMOD/021": "LVMOD/021 95mm2 1/2 length branch type connector module",
    "LVMOD/022": "LVMOD/022 300mm2 1/2 length branch type connector module",
    "LVMOD/023": "LVMOD/023 95mm2 Service Bridge Piece module",
    "LVMOD/024": "LVMOD/024 185mm2 Service Bridge Piece module",
    "LVMOD/025": "LVMOD/025 300mm2 Service Bridge Piece module",
    "LVMOD/026": "LVMOD/026 upto 35mm2 PILC service cable Earth Bond Kit module",
    "LVMOD/027": "LVMOD/027 50-95mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/028": "LVMOD/028 >95-185mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/029": "LVMOD/029 >185-300mm2 PILC Mains cable Earth Bond Kit module",
    "LVMOD/030": "LVMOD/030 Torque Limiting shear-off device module",
    "LVMOD/031": "LVMOD/031 95mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/032": "LVMOD/032 185mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/033": "LVMOD/033 300mm2 Aluminium mechanical shear-off lug module",
    "LVMOD/034": "LVMOD/034 480-740mm2 range taking Aluminium mechanical shear-off lug module",
    "LVMOD/035": "LVMOD/035 95mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/036": "LVMOD/036 185mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/037": "LVMOD/037 300mm2 Aluminium mechanical shear-off Busbar connector module",
    "LVMOD/038": "LVMOD/038 70-95mm2 pole termination module kit for 4c overhead lines and fuses",
    "LVMOD/039": "LVMOD/039 185mm pole termination module kit for 4c overhead lines and fuses",
    "LVMOD/040": "LVMOD/040 35-70mm2 Brass shear off lug module",
    "LVMOD/041": "LVMOD/041 60-120mm2 Brass shear off lug module"
}

hv_joint_termination_keys = {
    "11kv XLPE 3c Straight joint": "11kV XLPE 3c Straight Joint",
    "11kV 95mm XLPE trif joint": "11kV 95mm XLPE Trifurcating Joint",
    "11kV 185 - 300mm XLPE Trif joint": "11kV 185-300mm XLPE Trifurcating Joint",
    "11kV up to 70mm PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (up to 70mm)",
    "11kV 95-185 PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (95-185mm)",
    "11kV 185-300 PILC/PICAS to XLPE Joint": "11kV PILC/PICAS to XLPE Joint (185-300mm)",
    "11kV 95-185 XLPE to up to 70mm PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (95-185mm to 70mm)",
    "11kV 95-185 XLPE to 95-185 PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (95-185mm)",
    "11kV 185-300 XLPE to 185-300 PILC/PICAS Transition Trif Joint": "11kV XLPE to PILC/PICAS Transition Trif Joint (185-300mm)",
    "11kV Earthing kit for CORAL cables": "11kV Earthing Kit for CORAL Cables",
    "11kV Earthing kit for 50-300mm PILC cables": "11kV Earthing Kit for PILC Cables (50-300mm)",
    "11kV Earthing kit for up to 50mm PILC cables": "11kV Earthing Kit for PILC Cables (up to 50mm)",
    "11kV Build up kit for PILC / CORAL cables": "11kV Build Up Kit for PILC/CORAL Cables",
    "11kV Build up kit for XLPE cables": "11kV Build Up Kit for XLPE Cables",
    "11kV 95/185mm module for PAPER to PAPER joint": "11kV Paper to Paper Joint Module (95/185mm)",
    "11kV 300mm module for PAPER to PAPER joint": "11kV Paper to Paper Joint Module (300mm)",
    "11kV pole Term 1c 95mm": "11kV Pole Termination 1c 95mm",
    "11kV pole Term 1c 185/300mm": "11kV Pole Termination 1c 185/300mm",
    "11kV pole Term 3c 95mm": "11kV Pole Termination 3c 95mm",
    "11kV pole Term 3c 185/300mm": "11kV Pole Termination 3c 185/300mm",
    "OUTDR TERMN POLE STEELWORK 11 KV": "11kV Outdoor Pole Termination Steelwork",
    "11kV 95mm cable clamp for crucifix": "11kV Cable Clamp for Crucifix (95mm)",
    "11kV 185mm cable clamp for crucifix": "11kV Cable Clamp for Crucifix (185mm)",
    "11kV Surge Arrestor (Each)": "11kV Surge Arrestor",
    "33kv Joint Transition Trif (H-Type)": "33kV Joint Transition Trifurcating (H-Type)",
    "33kv Joint Trif (HSL-Type)": "33kV Joint Trifurcating (HSL-Type)",
    "33kv Joint 0.1 sq inch connectors (3 phases)": "33kV Joint Connectors 0.1 sq inch",
    "33kv Joint 0.4/0.5 sq inch connector (per phase)": "33kV Joint Connector 0.4/0.5 sq inch",
    "33kv Joint Connectors for Trif 150/300 Pilc": "33kV Joint Connectors for Trifurcating 150/300 PILC",
    "33kv Joint Straight up to 240mm (per phase)": "33kV Straight Joint (up to 240mm)",
    "33kv Joint Straight over 240mm needs connector (per phase)": "33kV Straight Joint (over 240mm)",
    "33kv Joint 400mm connector (each)": "33kV Joint Connector 400mm",
    "33kv Joint Transition 150/240mm to 0.3 PILC (per phase)": "33kV Joint Transition 150/240mm to 0.3 PILC",
    "11/33kv Pot End module up to 300mm (3 phases)": "11/33kV Pot End Module (up to 300mm)",
    "33kV Pole Term 1c 150-240mm (3 phase set)": "33kV Pole Termination 1c 150-240mm",
    "33kV Pole Term 1c 400-630mm (3 phase set)": "33kV Pole Termination 1c 400-630mm",
    "33kV Cable cleats for pole terms": "33kV Cable Cleats for Pole Terminations",
    "33kV Surge Arrestor 36kV (Each)": "33kV Surge Arrestor 36kV"
}

cable_accessory_keys = {
    "End cap up to 17mm diameter (25(1))": "End cap up to 17mm diameter (25(1))",
    "End cap 17-30mm dia(35(3))": "End cap 17-30mm dia(35(3))",
    "End Cap 30-45mm dia (95 LV or HV)": "End Cap 30-45mm dia (95 LV or HV)",
    "End Cap 45-95mm dia (185-300 LV or HV)": "End Cap 45-95mm dia (185-300 LV or HV)",
    "Ducting 32mm (OD 38mm) per metre (100m coil)": "Ducting 32mm (OD 38mm) per metre (100m coil)",
    "Ducting 50mm (OD 58mm) per metre (50m coil)": "Ducting 50mm (OD 58mm) per metre (50m coil)",
    "Ducting 100mm (3m Length) (90 in pallet)": "Ducting 100mm (3m Length) (90 in pallet)",
    "Ducting bend (100mm / 11.25 degree)": "Ducting bend (100mm / 11.25 degree)",
    "Ducting bend (100mm / 22.5 degree)": "Ducting bend (100mm / 22.5 degree)",
    "Ducting bend (100mm / 45 degree)": "Ducting bend (100mm / 45 degree)",
    "Ducting 150mm (3m Length) (39 in pallet)": "Ducting 150mm (3m Length) (39 in pallet)",
    "Ducting bend (150mm / 11.25 degree)": "Ducting bend (150mm / 11.25 degree)",
    "Ducting bend (150mm / 22.5 degree)": "Ducting bend (150mm / 22.5 degree)",
    "Ducting bend (150mm / 45 degree)": "Ducting bend (150mm / 45 degree)",
    "Resin 2 litre JEM Permanent": "Resin 2 litre JEM Permanent",
    "Resin 6 litre JEM Permanent": "Resin 6 litre JEM Permanent"
}

foundation_steelwork_keys = {
    "H' Pole Foundation Brace Steelwork for P6.010mm Centres ( Ref. SP4017651 )": "H' Pole Foundation Brace Steelwork for P6.010mm Centres ( Ref. SP4017651 )",
    "'H' Pole Foundation Brace Steelwork for 2500mm Centres ( Ref. SP4017652 )": "'H' Pole Foundation Brace Steelwork for 2500mm Centres ( Ref. SP4017652 )",
    "Stay / Foundation Block Type 1; 850mm as SP4019020": "Stay / Foundation Block Type 1; 850mm as SP4019020",
    "Stay / Foundation Block Type 2; 1300mm as SP4019020": "Stay / Foundation Block Type 2; 1300mm as SP4019020",
    "Foundation Block Type 3; 1500mm as SP4019020": "Foundation Block Type 3; 1500mm as SP4019020"
}

categories = [
    ("Poles 🪵", pole_keys, "Quantity"),
    ("Transformers ⚡🏭", transformer_keys, "Quantity"),
    ("Conductors", conductor_keys, "Length (Km)"),
    ("Conductors_2", conductor_2_keys, "Length (Km)"),
    ("Equipment", equipment_keys, "Quantity"),
    ("Insulators", insulator_keys, "Quantity"),
    ("LV Joints (Kits)", lv_joint_kit_keys, "Quantity"),
    ("LV Joint Modules", lv_joint_module_keys, "Quantity"),
    ("HV Joints / Terminations ⚡", hv_joint_termination_keys, "Quantity"),
    ("Cable Accessories 🔌", cable_accessory_keys, "Quantity"),
    ("Foundation & Steelwork 🏗️", foundation_steelwork_keys, "Quantity")
]


# ==============================================
# FUZZY MATCHING AND DATA MERGING FUNCTIONS
# ==============================================

def preprocess_string(text):
    """Preprocess string for better matching"""
    if pd.isna(text):
        return ""
    text = str(text).lower().strip()
    # Remove extra whitespace and special characters
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text

def find_best_match(target, candidates, threshold=75):
    """Find the best fuzzy match for a target string among candidates"""
    # Explicitly check for NaN values
    if pd.isna(target):
        return None, 0
    if candidates is None or len(candidates) == 0:
        return None, 0
    
    best_match = None
    best_score = 0
    
    for candidate in candidates:
        if pd.isna(candidate):
            continue
            
        # Calculate similarity score
        score = fuzz.partial_ratio(preprocess_string(target), preprocess_string(candidate))
        
        if score > best_score and score >= threshold:
            best_score = score
            best_match = candidate
    
    return best_match, best_score

@st.cache_data
def merge_with_pid_fuzzy(agg_df, pid_df):
    """
    Merge aggregated data with PID data using fuzzy matching on project descriptions.
    Requirements:
    1. Same project name
    2. Same shire
    3. Similar segment description (≥75% similarity)
    """
    if pid_df is None or agg_df.empty:
        return agg_df.copy()
    
    # Make copies to avoid modifying originals
    agg_df = agg_df.copy()
    pid_df = pid_df.copy()
    
    # Ensure required columns exist
    required_agg_cols = ['project', 'shire', 'segmentdesc']
    required_pid_cols = ['project', 'shire', 'project_description', 'pid_ohl_nr']
    
    missing_agg = [col for col in required_agg_cols if col not in agg_df.columns]
    missing_pid = [col for col in required_pid_cols if col not in pid_df.columns]
    
    if missing_agg:
        st.warning(f"Missing columns in aggregated data: {missing_agg}")
        return agg_df
    
    if missing_pid:
        st.warning(f"Missing columns in PID data: {missing_pid}")
        return agg_df
    
    # Create a list to store matched rows
    matched_rows = []
    
    # Get unique combinations from aggregated data
    unique_combinations = agg_df[['project', 'shire']].drop_duplicates()
    
    # Progress tracking
    progress_bar = st.progress(0)
    total_combinations = len(unique_combinations)
    
    st.sidebar.info(f"Fuzzy matching {total_combinations} project/shire combinations...")
    
    for idx, (project, shire) in enumerate(unique_combinations.itertuples(index=False, name=None)):
        # Update progress
        progress_bar.progress((idx + 1) / total_combinations)
        
        # Filter aggregated data for this project/shire combination
        agg_subset = agg_df[(agg_df['project'] == project) & (agg_df['shire'] == shire)]
        
        # Filter PID data for the same project/shire combination
        pid_subset = pid_df[(pid_df['project'] == project) & (pid_df['shire'] == shire)]
        
        if pid_subset.empty:
            # No PID data for this combination, keep original rows
            matched_rows.append(agg_subset)
            continue
        
        # Get unique segment descriptions from PID data
        pid_descriptions = pid_subset['project_description'].dropna().unique()
        
        # For each row in aggregated data, find best match
        for _, agg_row in agg_subset.iterrows():
            segment_desc = agg_row['segmentdesc']
            
            if pd.isna(segment_desc):
                # No segment description to match, keep as is
                matched_rows.append(pd.DataFrame([agg_row]))
                continue
            
            # Find best fuzzy match
            best_match, match_score = find_best_match(segment_desc, pid_descriptions, threshold=75)
            
            if best_match and match_score >= 75:
                # Get the PID OHL Nr for the matched description
                matched_pid_row = pid_subset[
                    pid_subset['project_description'] == best_match
                ].iloc[0]
                
                # Create a new row with PID information
                new_row = agg_row.copy()
                new_row['matched_project_description'] = best_match
                new_row['match_score'] = match_score
                new_row['pid_ohl_nr'] = matched_pid_row['pid_ohl_nr']
                new_row['project_description'] = best_match
                
                # Add other PID columns if they exist
                for col in pid_subset.columns:
                    if col not in new_row.index and col not in ['project', 'shire']:
                        new_row[col] = matched_pid_row[col]
                
                matched_rows.append(pd.DataFrame([new_row]))
            else:
                # No good match found, keep original
                matched_rows.append(pd.DataFrame([agg_row]))
    
    progress_bar.empty()
    
    # Combine all matched rows
    if matched_rows:
        result_df = pd.concat(matched_rows, ignore_index=True)
        
        # Fill NaN for unmatched rows
        result_df['matched_project_description'] = result_df['matched_project_description'].fillna('No Match')
        result_df['match_score'] = result_df['match_score'].fillna(0)
        result_df['pid_ohl_nr'] = result_df['pid_ohl_nr'].fillna('Not Available')
        
        # Count matches
        matched_count = (result_df['match_score'] >= 75).sum()
        total_count = len(result_df)
        
        st.sidebar.success(f"✓ Fuzzy matching complete: {matched_count}/{total_count} rows matched (≥75%)")
        
        return result_df
    else:
        return agg_df

# ==============================================
# MAIN APPLICATION
# ==============================================

# --- Gradient background ---
gradient_bg = """
<style>
    .stApp {
        background: linear-gradient(
            90deg,
            rgba(41, 28, 66, 1) 10%, 
            rgba(36, 57, 87, 1) 35%
        );
        color: white;
    }
</style>
"""
st.markdown(gradient_bg, unsafe_allow_html=True)

# --- Load logos ---
try:
    logo_left = Image.open("Images/GaeltecImage.png").resize((80, 80))
    logo_right = Image.open("Images/SPEN.png").resize((160, 80))
except:
    logo_left = logo_right = None
    st.warning("Logo images not found")

# --- Header layout ---
col1, col2, col3 = st.columns([1, 4, 1])
with col1:
    if logo_left:
        st.image(logo_left)
with col2:
    st.markdown("<h1 style='text-align:center; margin:0;'>Gaeltec Utilities.UK</h1>", unsafe_allow_html=True)
with col3:
    if logo_right:
        st.image(logo_right)

st.markdown("<h1>📊 Data Management Dashboard</h1>", unsafe_allow_html=True)

# ==============================================
# DATA LOADING AND PROCESSING
# ==============================================

# Load all data files
data = load_data_files()

# Check if main data is loaded
if data.get("aggregated") is None:
    st.error("Failed to load main aggregated data. Please check if 'CF_aggregated.parquet' exists.")
    st.stop()

agg_df = data["aggregated"]
resume_df = data["resume"]
misc_df = data["miscellaneous"]
pid_df = data["pid"]

# Process date columns in aggregated data
if "datetouse" in agg_df.columns:
    agg_df["datetouse_dt"] = pd.to_datetime(agg_df["datetouse"], errors="coerce")
    agg_df["datetouse_display"] = agg_df["datetouse_dt"].dt.strftime("%d/%m/%Y")
    agg_df.loc[agg_df["datetouse_dt"].isna(), "datetouse_display"] = "Unplanned"
    agg_df["datetouse_dt"] = agg_df["datetouse_dt"].dt.normalize()
else:
    agg_df["datetouse_dt"] = pd.NaT
    agg_df["datetouse_display"] = "Unplanned"

# ==============================================
# FUZZY MERGE WITH PID DATA
# ==============================================

st.sidebar.header("🔗 Data Integration")
fuzzy_match_enabled = st.sidebar.checkbox("Enable Fuzzy Matching with PID Data", value=True, 
                                          help="Match aggregated data with PID data using project descriptions (≥75% similarity)")

if fuzzy_match_enabled and pid_df is not None:
    with st.spinner("Performing fuzzy matching with PID data..."):
        enriched_df = merge_with_pid_fuzzy(agg_df, pid_df)
        
        # Show matching statistics
        if 'match_score' in enriched_df.columns:
            matched = (enriched_df['match_score'] >= 75).sum()
            total = len(enriched_df)
            st.sidebar.info(f"**Fuzzy Match Results:**\n{matched}/{total} rows matched ({(matched/total*100):.1f}%)")
else:
    enriched_df = agg_df.copy()
    if fuzzy_match_enabled and pid_df is None:
        st.sidebar.warning("PID data not available for fuzzy matching")

# ==============================================
# APPLY MISC DATA MAPPING
# ==============================================

if misc_df is not None and "item" in enriched_df.columns:
    # Map material codes from misc data
    misc_df["column_b"] = misc_df["column_b"].astype(str)
    material_map = misc_df.set_index("column_b")["column_k"].to_dict()
    enriched_df["material_code"] = enriched_df["item"].map(material_map)
else:
    enriched_df["material_code"] = None

# ==============================================
# SIDEBAR FILTERS
# ==============================================

st.sidebar.header("📊 Filters")

def multi_select_filter(col, label, df, parent=None):
    """Create multi-select filter"""
    if col not in df.columns:
        return ["All"], df
    
    temp = df.copy()
    if parent and "All" not in parent[1]:
        temp = temp[temp[parent[0]].isin(parent[1])]
    
    options = ["All"] + sorted(temp[col].dropna().unique())
    selected = st.sidebar.multiselect(label, options, default=["All"])
    
    if "All" not in selected:
        temp = temp[temp[col].isin(selected)]
    
    return selected, temp

# Initialize filtered dataframe
filtered_df = enriched_df.copy()

# Apply filters sequentially
selected_shire, filtered_df = multi_select_filter("shire", "Shire", filtered_df)
selected_project, filtered_df = multi_select_filter(
    "project", "Project", filtered_df, ("shire", selected_shire)
)
selected_pm, filtered_df = multi_select_filter(
    "projectmanager", "Project Manager", filtered_df, ("shire", selected_shire)
)

# Add PID OHL Nr filter if available
if 'pid_ohl_nr' in filtered_df.columns:
    selected_pid, filtered_df = multi_select_filter(
        "pid_ohl_nr", "PID OHL Nr", filtered_df
    )

# ==============================================
# DATE FILTER
# ==============================================

st.sidebar.markdown("---")
st.sidebar.subheader("📅 Date Range Filter")

date_range_str = "All dates"

if 'datetouse_dt' in filtered_df.columns:
    min_date = filtered_df['datetouse_dt'].min()
    max_date = filtered_df['datetouse_dt'].max()
    
    if pd.notna(min_date) and pd.notna(max_date):
        date_range = st.sidebar.date_input(
            "Select date range:",
            value=(min_date.date(), max_date.date()),
            min_value=min_date.date(),
            max_value=max_date.date()
        )
        
        if len(date_range) == 2:
            start_date, end_date = date_range
            start_dt = pd.Timestamp(start_date)
            end_dt = pd.Timestamp(end_date) + pd.Timedelta(days=1) - pd.Timedelta(seconds=1)
            
            filtered_df = filtered_df[
                (filtered_df['datetouse_dt'] >= start_dt) &
                (filtered_df['datetouse_dt'] <= end_dt)
            ]
            date_range_str = f"{start_date.strftime('%d/%m/%Y')} to {end_date.strftime('%d/%m/%Y')}"
    else:
        date_range_str = "Date data unavailable"
        st.sidebar.info("No valid date data available")
else:
    date_range_str = "No date column"
    st.sidebar.info("Date column not found in data")

# ==============================================
# DASHBOARD DISPLAY - FINANCIAL SUMMARY
# ==============================================

# Calculate totals
total_sum, variation_sum = 0, 0
if 'total' in filtered_df.columns:
    total_series = pd.to_numeric(
        filtered_df['total'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
        errors='coerce'
    )
    total_sum = total_series.sum(skipna=True)
    
    if 'orig' in filtered_df.columns:
        orig_series = pd.to_numeric(
            filtered_df['orig'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
            errors='coerce'
        )
        variation_sum = (total_series - orig_series).sum(skipna=True)

formatted_total = f"{total_sum:,.2f}".replace(",", " ").replace(".", ",")
formatted_variation = f"{variation_sum:,.2f}".replace(",", " ").replace(".", ",")

# Display Total & Variation
st.markdown("<h2>Financial</h2>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align:center; color:white;'>Revenue</h3>", unsafe_allow_html=True)

# Try to display money logo
try:
    money_logo = Image.open("Images/Pound.png").resize((40, 40))
    buffered = BytesIO()
    money_logo.save(buffered, format="PNG")
    money_logo_base64 = base64.b64encode(buffered.getvalue()).decode()
    
    has_logo = True
except:
    has_logo = False

if has_logo:
    st.markdown(
        f"""
        <div style='display:flex; justify-content:center;'>
            <div style='display:flex; flex-direction:column; gap:4px;'>
                <div style='display:flex; align-items:center; gap:10px;'>
                    <h2 style='color:#32CD32; margin:0; font-size:36px;'><b>Total:</b> {formatted_total}</h2>
                    <img src='data:image/png;base64,{money_logo_base64}' width='40' height='40'/>
                </div>
                <div style='display:flex; align-items:center; gap:8px;'>
                    <h2 style='color:#32CD32; font-size:25px; margin:0;'><b>Variation:</b> {formatted_variation}</h2>
                    <img src='data:image/png;base64,{money_logo_base64}' width='28' height='28'/>
                </div>
                <p style='text-align:center; font-size:14px; margin-top:4px;'>
                    ({date_range_str} | Shires: {len(selected_shire) if 'All' not in selected_shire else 'All'} | 
                    Projects: {len(selected_project) if 'All' not in selected_project else 'All'} | 
                    PMs: {len(selected_pm) if 'All' not in selected_pm else 'All'})
                </p>
            </div>
        </div>
        """,
        unsafe_allow_html=True
    )
else:
    st.markdown(
        f"""
        <div style='text-align:center;'>
            <h2 style='color:#32CD32;'>Total: {formatted_total}</h2>
            <h3 style='color:#32CD32;'>Variation: {formatted_variation}</h3>
            <p style='font-size:14px;'>
                ({date_range_str} | Shires: {len(selected_shire) if 'All' not in selected_shire else 'All'} | 
                Projects: {len(selected_project) if 'All' not in selected_project else 'All'})
            </p>
        </div>
        """,
        unsafe_allow_html=True
    )

# ==============================================
# REVENUE CHART
# ==============================================

st.markdown("<h3 style='text-align:center; color:white;'>Revenue Trend</h3>", unsafe_allow_html=True)

try:
    if not filtered_df.empty and 'total' in filtered_df.columns and 'datetouse_dt' in filtered_df.columns:
        chart_df = filtered_df[filtered_df['datetouse_dt'].notna()].copy()
        chart_df = chart_df[chart_df['datetouse_dt'] >= '2000-01-01']
        chart_df['total'] = pd.to_numeric(chart_df['total'], errors='coerce')
        chart_df = chart_df[chart_df['total'].notna()]
        
        if not chart_df.empty:
            revenue_by_date = chart_df.groupby('datetouse_dt')['total'].sum().reset_index()
            revenue_by_date = revenue_by_date.sort_values('datetouse_dt')
            
            fig_revenue = px.line(
                revenue_by_date,
                x='datetouse_dt',
                y='total',
                title="Daily Revenue",
                labels={'datetouse_dt': 'Date', 'total': 'Revenue (£)'}
            )
            
            fig_revenue.update_traces(
                mode='lines+markers',
                line=dict(width=3, color='#32CD32'),
                marker=dict(size=6, color='#32CD32'),
                hovertemplate='<b>Date: %{x}</b><br>Revenue: £%{y:,.0f}<extra></extra>'
            )
            
            fig_revenue.update_layout(
                height=500,
                xaxis=dict(
                    tickformat="%d %b %Y",
                    tickangle=45,
                    gridcolor='rgba(128,128,128,0.2)',
                    rangeslider=dict(visible=True),
                    type='date'
                ),
                yaxis=dict(
                    title='Revenue (£)',
                    tickformat=",.0f",
                    gridcolor='rgba(128,128,128,0.2)'
                ),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                font=dict(color='white'),
                title_font_size=16
            )
            
            st.plotly_chart(fig_revenue, use_container_width=True)
        else:
            st.info("No projects with valid dates and totals for selected filters.")
    else:
        st.info("No data available for revenue chart.")
except Exception as e:
    st.warning(f"Could not generate revenue chart: {str(e)}")

# ==============================================
# PROJECTS AND WORKS SECTIONS (Adapted for enriched data)
# ==============================================

# Display Projects Distribution
col_top_left, col_top_right = st.columns([1, 1])

with col_top_left:
    st.markdown("<h3 style='text-align:center; color:white;'>Projects Distribution</h3>", unsafe_allow_html=True)
    
    try:
        if not filtered_df.empty and 'project' in filtered_df.columns:
            project_counts = filtered_df['project'].value_counts().reset_index()
            project_counts.columns = ['Project', 'total']
            
            if len(project_counts) > 8:
                top_projects = project_counts.head(7)
                other_count = project_counts['total'].iloc[7:].sum()
                other_row = pd.DataFrame({'Project': ['Other'], 'total': [other_count]})
                project_data = pd.concat([top_projects, other_row], ignore_index=True)
            else:
                project_data = project_counts
            
            fig_projects = px.pie(
                project_data,
                names='Project',
                values='total',
                title="",
                hole=0.4
            )
            fig_projects.update_traces(
                textinfo='percent+label',
                textfont_size=14,
                marker=dict(line=dict(color='#000000', width=1))
            )
            fig_projects.update_layout(
                title_text="",
                title_font_size=16,
                font=dict(color='white'),
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                showlegend=False,
                annotations=[dict(text=f'Total<br>{len(filtered_df)}', x=0.5, y=0.5, font_size=16, showarrow=False)]
            )
            
            st.plotly_chart(fig_projects, use_container_width=True)
        else:
            st.info("No project data available.")
    except Exception as e:
        st.warning(f"Could not generate projects pie chart: {str(e)}")

# Works total
with col_top_right:
    # Left side: Projects & Segments Overview and Works Complete pie chart
    col_left_top, col_left_bottom = st.columns([1, 1])
    
    with col_left_top:
        st.markdown("<h3 style='color:white;'>Projects & Segments Overview</h3>", unsafe_allow_html=True)

        if 'project' in filtered_df.columns and 'segmentcode' in filtered_df.columns:
            projects = filtered_df['project'].dropna().unique()
            if len(projects) == 0:
                st.info("No projects found for the selected filters.")
            else:
                for proj in sorted(projects):
                    segments = filtered_df[filtered_df['project'] == proj]['segmentcode'].dropna().unique()
                
                    # Use expander to make segment list scrollable
                    with st.expander(f"Project: {proj} ({len(segments)} segments)"):
                        if len(segments) > 0:
                            # Scrollable container for segments
                            st.markdown(
                                "<div style='max-height:150px; overflow-y:auto; padding:5px; border:1px solid #444;'>"
                                + "<br>".join(segments.astype(str))
                                + "</div>",
                                unsafe_allow_html=True
                            )
                        else:
                            st.write("No segment codes for this project.")
        else:
            st.info("Project or Segment Code columns not found in the data.")

# -------------------------------
# --- Works Complete Pie Chart ---
# -------------------------------
st.markdown("<h3 style='text-align:center; color:white;'>Works Complete</h3>", unsafe_allow_html=True)
try:
    if 'resume_df' in locals():
        filtered_segments = filtered_df['segment'].dropna().astype(str).str.strip().str.lower().unique()
        resume_df['section'] = resume_df['section'].dropna().astype(str).str.strip().str.lower()

        if {'section', '%complete'}.issubset(resume_df.columns):
            resume_filtered = resume_df[resume_df['section'].isin(filtered_segments)]

            if not resume_filtered.empty:
                avg_complete = resume_filtered['%complete'].mean()
                avg_complete = min(max(avg_complete, 0), 100)

                pie_data = pd.DataFrame({
                    'Status': ['Completed', 'Done or Remaining'],
                    'Value': [avg_complete, 100 - avg_complete]
                })

                fig_pie = px.pie(
                    pie_data,
                    names='Status',
                    values='Value',
                    color='Status',
                    color_discrete_map={'Completed': 'green', 'Done or Remaining': 'red'},
                    hole=0.6
                )
                fig_pie.update_traces(
                    textinfo='percent+label',
                    textfont_size=20
                )
                fig_pie.update_layout(
                    title_text="",
                    title_font_size=20,
                    font=dict(color='white'),
                    paper_bgcolor='rgba(0,0,0,0)',
                    plot_bgcolor='rgba(0,0,0,0)',
                    showlegend=True,
                    legend=dict(font=dict(color='white'))
                )

                st.plotly_chart(fig_pie, use_container_width=True)
            else:
                st.info("No matching sections found for the selected filters to generate % completion chart.")

except Exception as e:
    st.warning(f"Could not generate % Complete pie chart: {e}")
    
# -------------------------------
# --- Map Section ---
# -------------------------------
col_map, col_desc = st.columns([2, 1])
with col_map:
    st.header("🗺️ Regional Map View")
    folder_path = r"Maps"
    file_list = glob.glob(os.path.join(folder_path, "*.json"))

    if not file_list:
        st.error(f"No JSON files found in folder: {folder_path}")
    else:
        gdf_list = [gpd.read_file(file) for file in file_list]
        combined_gdf = gpd.GeoDataFrame(pd.concat(gdf_list, ignore_index=True), crs=gdf_list[0].crs)

        if "region" in filtered_df.columns:
            active_regions = filtered_df["region"].dropna().unique().tolist()
            wards_to_select = []
            for region in active_regions:
                if region in mapping_region:
                    wards_to_select.extend(mapping_region[region])
                else:
                    wards_to_select.append(region)
            wards_to_select = list(set(wards_to_select))
            areas_of_interest = combined_gdf[combined_gdf["WD13NM"].isin(wards_to_select)]
        else:
            areas_of_interest = pd.DataFrame()

        if not areas_of_interest.empty:
            areas_of_interest["geometry_simplified"] = areas_of_interest.geometry.simplify(tolerance=0.01)
            centroid = areas_of_interest.geometry_simplified.centroid.unary_union.centroid

            # Red flag
            flag_data = pd.DataFrame({"lon": [centroid.x], "lat": [centroid.y], "icon_name": ["red_flag"]})
            icon_mapping = {
                "red_flag": {
                    "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/3/3e/Red_flag_icon.svg/128px-Red_flag_icon.png",
                    "width": 128, "height": 128, "anchorY": 128
                }
            }

            polygon_layer = pdk.Layer(
                "GeoJsonLayer",
                areas_of_interest["geometry_simplified"].__geo_interface__,
                stroked=True,
                filled=True,
                get_fill_color=[160, 120, 80, 200],
                get_line_color=[0, 0, 0],
                pickable=True
            )

            flag_layer = pdk.Layer(
                "IconLayer",
                data=flag_data,
                get_icon="icon_name",
                get_size=4,
                size_scale=15,
                get_position='[lon, lat]',
                pickable=True,
                icon_mapping=icon_mapping
            )

            view_state = pdk.ViewState(latitude=centroid.y, longitude=centroid.x, zoom=8, pitch=0)

            st.pydeck_chart(
                pdk.Deck(
                    layers=[polygon_layer, flag_layer],
                    initial_view_state=view_state,
                    map_style="mapbox://styles/mapbox/outdoors-v11"
                )
            )
        else:
            st.info("No matching regions found for the selected filters.")


with col_desc:
    st.markdown("<h3 style='color:white;'>Weather</h3>", unsafe_allow_html=True)
    
    # --- Scottish Weather Widget ---
    try:
        # Get API key from secrets
        api_key = st.secrets.get("d4d09fcf1373f72c30b970fb20d51fd9")
        
        if not api_key:
            st.info("Weather API key not configured")
        else:
            # Location selector
            location = st.selectbox(
                "Select Location",
                ["Ayrshire", "Lanarkshire", "Glasgow", "Edinburgh"],
                index=0,
                key="weather_location"
            )
            
            if st.button("Refresh Weather", key="refresh_weather"):
                st.rerun()
            
            # Get current weather
            weather_data = get_scottish_weather(api_key, location)
            
            if weather_data:
                # Display weather information
                temp = weather_data['main']['temp']
                feels_like = weather_data['main']['feels_like']
                humidity = weather_data['main']['humidity']
                wind_speed = weather_data['wind']['speed']
                description = weather_data['weather'][0]['description'].title()
                icon_code = weather_data['weather'][0]['icon']
                
                # Weather icon and description
                col_icon, col_desc = st.columns([1, 2])
                with col_icon:
                    st.image(f"http://openweathermap.org/img/wn/{icon_code}@2x.png", width=50)
                with col_desc:
                    st.write(f"**{description}**")
                
                # Weather metrics
                st.metric("Temperature", f"{temp}°C", f"Feels like {feels_like}°C")
                st.metric("Humidity", f"{humidity}%")
                st.metric("Wind Speed", f"{wind_speed} m/s")
                
                # Construction impact assessment
                st.markdown("---")
                st.markdown("**Construction Impact:**")
                impact = assess_construction_impact(weather_data)
                st.write(impact)
            else:
                st.error("Failed to fetch weather data")
                
    except Exception as e:
        st.warning(f"Could not load weather information: {e}")


# -------------------------------
# --- Mapping Bar Charts + Drill-down + Excel Export ---
# -------------------------------
st.header("🪵 Materials")
convert_to_miles = st.checkbox("Convert Equipment/Conductor Length to Miles")

categories = [
    ("Poles 🪵", pole_keys, "Quantity"),
    ("Transformers ⚡🏭", transformer_keys, "Quantity"),
    ("Conductors", conductor_keys, "Length (Km)"),
    ("Conductors_2", conductor_2_keys, "Length (Km)"),
    ("Equipment", equipment_keys, "Quantity"),
    ("Insulators", insulator_keys, "Quantity"),
    ("LV Joints (Kits)", lv_joint_kit_keys, "Quantity"),
    ("LV Joint Modules", lv_joint_module_keys, "Quantity"),
    ("HV Joints / Terminations ⚡", hv_joint_termination_keys, "Quantity"),
    ("Cable Accessories 🔌", cable_accessory_keys, "Quantity"),
    ("Foundation & Steelwork 🏗️", foundation_steelwork_keys, "Quantity")
]

def sanitize_sheet_name(name: str) -> str:
    name = str(name)
    name = re.sub(r'[:\\/*?\[\]\n\r]', '_', name)
    name = re.sub(r'[^\x00-\x7F]', '_', name)  # remove Unicode like m²
    return name[:31]


for cat_name, keys, y_label in categories:

    # Only process if columns exist
    if 'item' not in filtered_df.columns or 'mapped' not in filtered_df.columns:
        st.warning("Missing required columns: item / mapped")
        continue
        
    # Build regex pattern for this category's keys
    pattern = '|'.join([re.escape(k) for k in keys.keys()])

    mask = filtered_df['item'].astype(str).str.contains(pattern, case=False, na=False)
    sub_df = filtered_df[mask]

    if sub_df.empty:
        st.info(f"No data found for {cat_name}")
        continue

    # Aggregate
    if 'qsub' in sub_df.columns:
        sub_df['qsub_clean'] = pd.to_numeric(
            sub_df['qsub'].astype(str).str.replace(" ", "").str.replace(",", ".", regex=False),
            errors='coerce'
        )
        bar_data = sub_df.groupby('mapped')['qsub_clean'].sum().reset_index()
        bar_data.columns = ['Mapped', 'Total']
    else:
        bar_data = sub_df['mapped'].value_counts().reset_index()
        bar_data.columns = ['Mapped', 'Total']

    # Divide Conductors_2 by 1000
    if cat_name == "Conductors_2":
        bar_data['Total'] = bar_data['Total'] / 1000

    # Convert conductor units if needed
    y_axis_label = y_label
    if cat_name in ["Conductors", "Conductors_2"] and convert_to_miles:
        bar_data['Total'] = bar_data['Total'] * 0.621371
        y_axis_label = "Length (Miles)"

    # Compute grand total for the category
    grand_total = bar_data['Total'].sum()

    # Update Streamlit subheader with total
    st.subheader(f"🔹 {cat_name} — Total: {grand_total:,.2f}")

    # Draw the bar chart
    fig = go.Figure(data=[
        go.Bar(
            x=bar_data['Mapped'].astype(str).tolist(),
            y=bar_data['Total'].astype(float).tolist(),
            text=bar_data['Total'].astype(float).tolist(),
            texttemplate='%{y:,.1f}',
            textposition='outside'
        )
    ])

    fig.update_layout(
        title=f"{cat_name} Overview",
        xaxis_title="Mapping",
        yaxis_title=y_axis_label
    )
    
    # Add background colors separately
    fig.update_layout(
        plot_bgcolor='rgba(0,0,0,0)',
        paper_bgcolor='rgba(0,0,0,0)',
        yaxis=dict(
            gridcolor='rgba(255,255,255,0.3)'  # Semi-transparent white grid
        )
    )

    # Display the chart
    st.plotly_chart(fig, use_container_width=True, height=500)

    # COLLAPSIBLE BUTTONS SECTION
    with st.expander("🔍 Click to explore more information", expanded=False):
        st.subheader("Select Mapping to Drill-down:")
        
        # Option 1: Buttons in columns
        cols = st.columns(3)  # 3 buttons per row
        
        for idx, mapping_value in enumerate(bar_data['Mapped']):
            col_idx = idx % 3  # Which column to use (0, 1, or 2)
            
            with cols[col_idx]:
                button_key = f"btn_{cat_name}_{mapping_value}_{idx}"
                
                if st.button(f"📊 {mapping_value}", key=button_key, use_container_width=True):
                    st.session_state[f"selected_{cat_name}"] = mapping_value
                    st.rerun()  # Refresh to show the details immediately

    # Check if a mapping was selected
    selected_mapping = st.session_state.get(f"selected_{cat_name}")
    
    if selected_mapping:
        st.subheader(f"Details for: **{selected_mapping}**")
        
        # Add a button to clear the selection
        if st.button("❌ Clear Selection", key=f"clear_{cat_name}"):
            del st.session_state[f"selected_{cat_name}"]
            st.rerun()
        
        selected_rows = sub_df[sub_df['mapped'] == selected_mapping].copy()
        selected_rows.columns = selected_rows.columns.str.strip().str.lower()
        selected_rows = selected_rows.loc[:, ~selected_rows.columns.duplicated()]

        if 'datetouse' in selected_rows.columns:
            selected_rows['datetouse_display'] = pd.to_datetime(
                selected_rows['datetouse'], errors='coerce'
            ).dt.strftime("%d/%m/%Y")
            selected_rows.loc[selected_rows['datetouse'].isna(), 'datetouse_display'] = "Unplanned"


        # Your original approach but working:
        extra_cols = ['poling team','team_name','segmentdesc','segmentcode', 'projectmanager', 'project', 'shire','material code' , 'sourcefile','pid_ohl_nr']

        # --- Rename columns for display ---
        rename_map = {
            "poling team": "code", 
            "team_name": "team lider"
        }
        selected_rows = selected_rows.rename(columns=rename_map)

        # Update extra_cols to match renamed columns
        for old, new in rename_map.items():
            extra_cols = [new if c == old else c for c in extra_cols]

        # Ensure all extra_cols exist in selected_rows
        for col in extra_cols:
            if col not in selected_rows.columns:
                selected_rows[col] = None

        # --- Create display date column ---
        if 'datetouse' in selected_rows.columns:
            selected_rows['datetouse_display'] = pd.to_datetime(
                selected_rows['datetouse'], errors='coerce'
            ).dt.strftime("%d/%m/%Y")
            selected_rows.loc[selected_rows['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

        # --- Columns to display ---
        display_cols = ['mapped','pole','qsub','datetouse_display'] + extra_cols
        display_cols = [c for c in display_cols if c in selected_rows.columns]

        # --- Display dataframe ---
        if not selected_rows.empty:
            st.dataframe(selected_rows[display_cols], use_container_width=True)
            st.write(f"**Total records:** {len(selected_rows)}")

            if 'qsub_clean' in selected_rows.columns:
                total_qsub = selected_rows['qsub_clean'].sum()
                st.write(f"Total QSUB: {total_qsub:,.2f}")
        else:
            st.info("No records found for this selection")

        # --- Excel Export: Aggregated ---
        buffer_agg = BytesIO()
        with pd.ExcelWriter(buffer_agg, engine='openpyxl') as writer:
            aggregated_df = pd.DataFrame()
            for bar_value in bar_data['Mapped']:
                df_bar = sub_df[sub_df['mapped'] == bar_value].copy()
                df_bar = df_bar.loc[:, ~df_bar.columns.duplicated()]
                if 'datetouse' in df_bar.columns:
                    df_bar['datetouse_display'] = pd.to_datetime(
                        df_bar['datetouse'], errors='coerce'
                    ).dt.strftime("%d/%m/%Y")
                    df_bar.loc[df_bar['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

                cols_to_include = ['mapped', 'datetouse_display'] + extra_cols
                cols_to_include = [c for c in cols_to_include if c in df_bar.columns]
                df_bar = df_bar[cols_to_include]

                aggregated_df = pd.concat([aggregated_df, df_bar], ignore_index=True)

            aggregated_df.to_excel(writer, sheet_name='Aggregated', index=False)

        buffer_agg.seek(0)
        st.download_button(
            f"📥 Download Excel (Aggregated): {cat_name} Details",
            buffer_agg,
            file_name=f"{cat_name}_Details_Aggregated.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Excel Export - Separate Sheets
        buffer_sep = BytesIO()
        with pd.ExcelWriter(buffer_sep, engine='openpyxl') as writer:
            for bar_value in bar_data['Mapped']:
                df_bar = sub_df[sub_df['mapped'] == bar_value].copy()
                df_bar = df_bar.loc[:, ~df_bar.columns.duplicated()]
                if 'datetouse' in df_bar.columns:
                    df_bar['datetouse_display'] = pd.to_datetime(
                        df_bar['datetouse'], errors='coerce'
                    ).dt.strftime("%d/%m/%Y")
                    df_bar.loc[df_bar['datetouse'].isna(), 'datetouse_display'] = "Unplanned"

                cols_to_include = ['mapped', 'datetouse_display'] + extra_cols
                cols_to_include = [c for c in cols_to_include if c in df_bar.columns]
                df_bar = df_bar[cols_to_include]

                sheet_name = sanitize_sheet_name(bar_value)
                df_bar.to_excel(writer, sheet_name=sheet_name, index=False)

        buffer_sep.seek(0)
        st.download_button(
            f"📥 Download Excel (Separated): {cat_name} Details",
            buffer_sep,
            file_name=f"{cat_name}_Details_Separated.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
# -----------------------------
# 🛠️ Works Section
# -----------------------------
st.header("🛠️ Works")

if misc_df is not None:
    # -----------------------------
    # Data preparation
    # -----------------------------
    filtered_df['item'] = filtered_df['item'].astype(str)
    misc_df['column_b'] = misc_df['column_b'].astype(str)

    # Map items to work instructions
    item_to_column_i = misc_df.set_index('column_b')['column_i'].to_dict()
    poles_df = filtered_df[filtered_df['pole'].notna() & (filtered_df['pole'].astype(str).str.lower() != "nan")].copy()
    poles_df['Work instructions'] = poles_df['item'].map(item_to_column_i)

    # Keep only rows with valid instructions, comments, and team_name
    poles_df_clean = poles_df.dropna(subset=['Work instructions', 'comment', 'team_name'])[
        ['pole', 'segmentcode', 'Work instructions', 'comment', 'team_name']
    ]

    # -----------------------------
    # 🔘 Segment selector
    # -----------------------------
    segment_options = ['All'] + sorted(poles_df_clean['segmentcode'].dropna().astype(str).unique())
    selected_segment = st.selectbox("Select a segment code:", segment_options)

    if selected_segment != 'All':
        poles_df_view = poles_df_clean[poles_df_clean['segmentcode'].astype(str) == selected_segment]
    else:
        poles_df_view = poles_df_clean.copy()

    # -----------------------------
    # 🎯 Pole selector (Cascading)
    # -----------------------------
    pole_options = sorted(poles_df_view['pole'].dropna().astype(str).unique())
    selected_pole = st.selectbox("Select a pole to view details:", ["All"] + pole_options)

    # Filter by selected pole
    if selected_pole != "All":
        poles_df_view = poles_df_view[poles_df_view['pole'].astype(str) == selected_pole]

    # Display pole details if one is selected
    if selected_pole != "All" and not poles_df_view.empty:
        st.write(f"Details for pole **{selected_pole}**:")
        st.dataframe(poles_df_view)

    # -----------------------------
    # 📊 Pie chart (Works breakdown)
    # -----------------------------

    if not poles_df_view.empty:
        # Count work instructions and remove NaN / empty strings
        work_data = (
            poles_df_view['Work instructions']
            .astype(str)
            .str.lower()
            .replace('nan', pd.NA)
            .dropna()  # remove NaN
            .value_counts()
            .reset_index()
        )
        work_data.columns = ['Work instructions', 'total']

        if not work_data.empty:
            fig_work = px.pie(
                work_data,
                names='Work instructions',
                values='total',
                hole=0.4
            )
            fig_work.update_traces(textinfo='percent+label', textfont_size=16)
            fig_work.update_layout(paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', showlegend=False)
            st.plotly_chart(fig_work, use_container_width=True)
        else:
            st.info("No valid work instructions available for the selected filters.")
    # -----------------------------
    # 📄 Word export
    # -----------------------------
    if not poles_df_view.empty:
        word_file = poles_to_word(poles_df_view)
        st.download_button(
            label="⬇️ Download Work Instructions (.docx)",
            data=word_file,
            file_name="Pole_Work_Instructions.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# -----------------------------
# 📈 Jobs per Team per Day (Segment + Pole aware)
# -----------------------------
st.subheader("📈 Jobs per Team per Day")

if agg_view is not None and 'total' in agg_view.columns:
    filtered_agg = agg_view.copy()

    # Apply segment filter
    if selected_segment != 'All' and 'segmentcode' in filtered_agg.columns:
        filtered_agg = filtered_agg[
            filtered_agg['segmentcode'].astype(str).str.strip() == str(selected_segment).strip()
        ]

    # Apply pole filter
    if selected_pole != "All" and 'pole' in filtered_agg.columns:
        filtered_agg = filtered_agg[
            filtered_agg['pole'].astype(str).str.strip() == str(selected_pole).strip()
        ]

    # Ensure datetime column
    if 'datetouse_dt' not in filtered_agg.columns:
        filtered_agg['datetouse_dt'] = pd.to_datetime(filtered_agg['datetouse'], errors='coerce')
    else:
        filtered_agg['datetouse_dt'] = pd.to_datetime(filtered_agg['datetouse_dt'], errors='coerce')

    # Ensure 'total' is numeric
    filtered_agg['total'] = pd.to_numeric(filtered_agg['total'], errors='coerce').fillna(0)

    # Drop invalid rows
    filtered_agg = filtered_agg.dropna(subset=['datetouse_dt', 'team_name'])

    if not filtered_agg.empty:
        # Aggregate per day per team
        time_df = filtered_agg.groupby(['datetouse_dt', 'team_name'], as_index=False)['total'].sum()

        # Plot line chart
        fig_time = px.line(
            time_df,
            x='datetouse_dt',
            y='total',
            color='team_name',
            markers=True,
            hover_data={'datetouse_dt': True, 'team_name': True, 'total': True}
        )
        fig_time.update_layout(
            xaxis_title="Day",
            yaxis_title="Total Jobs £",
            xaxis=dict(
                tickformat="%d/%m/%Y",
                tickangle=45,
                nticks=10,
                tickmode='auto',
            ),
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            legend_title_text="Team",
            height=500
        )
        st.plotly_chart(fig_time, use_container_width=True)
    else:
        st.info("No time-based data available for the selected filters.")
else:
    st.info("No 'total' column found in aggregated data.")
